import re
import os
import docx
import pandas as pd
import customtkinter as ctk
import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document

# =============================================================================
# ПАРАМЕТРЫ
# =============================================================================

TABLE_NUMBER = 2   # Какая по счёту таблица (1=первая, 2=вторая и т.д.)
START_ROW = 3      # С какой строки таблицы читать

# Удаляем только col1 и col7, а col5,col6 оставляем (будем обрабатывать):
EXCLUDED_COLS = ['col1', 'col7']

# Переименуем служебные колонки
RENAME_MAP = {
    'semester':   'Семестр',
    'discipline': 'Дисциплина',
    'col2':       'Тип занятия',
    'col3':       'Время в минутах',
    'col4':       'Учебные вопросы',
    'col5':       'Материальное обеспечение на занятие',
    'col6':       'Литература на занятие',
}

# Итоговый порядок столбцов
DESIRED_ORDER = [
    'Дисциплина',
    'Название темы',
    'Номер темы',
    'Тип занятия',
    'Название занятия',
    'Номер занятия',
    'Учебные вопросы',
    'Время в минутах',
    'Материальное обеспечение на занятие',
    'Литература на занятие',
    'Знать',
    'Уметь',
    'Владеть',
    #Доп поля
    'Семестр',
    'Литература',
    'Материальное обеспечение',
]

# -----------------------------------------------------------------------------
# КОНСТАНТЫ: маркеры для литературы, мат. обеспечения и т.д.
# -----------------------------------------------------------------------------

START_REF = "III. ЛИТЕРАТУРА"
END_REF   = "Материальное обеспечение занятия"

START_MATERIAL = "Материальное обеспечение занятия"
END_MATERIAL   = "IV."

KNOW_START   = "Знать:"
KNOW_END     = "Уметь:"
SKILL_START  = "Уметь:"
SKILL_END    = "Владеть:"
MASTER_START = "Владеть:"
MASTER_END   = "Основными видами занятий по дисциплине"

# =============================================================================
# ФУНКЦИИ ДЛЯ ПОИСКА ТЕКСТА МЕЖДУ МАРКЕРАМИ
# =============================================================================

def get_discipline_name(doc):
    """
    Ищет строку вида: 'изучения дисциплины «XXX»' в документе,
    возвращает XXX или None, если не найдено.
    """
    pattern = r'изучения\s+дисциплины\s+«([^»]+)»'
    for paragraph in doc.paragraphs:
        txt = paragraph.text.strip()
        match = re.search(pattern, txt, flags=re.IGNORECASE)
        if match:
            return match.group(1)
    return None

def parse_text_between(doc, start_text, end_text):
    """
    Собирает параграфы между start_text и end_text (регистронезависимо),
    возвращая многострочную строку (разделитель '\n').
    Если end_text не найден, берем до конца документа.
    """
    start_found = False
    results = []
    start_up = start_text.upper()
    end_up = end_text.upper()

    for paragraph in doc.paragraphs:
        txt = paragraph.text.strip()
        up = txt.upper()

        if not start_found and up.startswith(start_up):
            start_found = True
            continue

        if start_found:
            if up.startswith(end_up):
                break
            if txt:
                results.append(txt)

    return "\n".join(results)

# =============================================================================
# ЧТЕНИЕ ТАБЛИЦЫ
# =============================================================================

def read_table_from_docx(doc, table_num, start_row):
    all_tables = doc.tables
    if len(all_tables) < table_num:
        raise ValueError(f"В документе {len(all_tables)} таблиц, а запрошена №{table_num}.")

    table = all_tables[table_num - 1]
    rows = table.rows

    if len(rows) < start_row:
        raise ValueError(f"В таблице №{table_num} всего {len(rows)} строк, "
                         f"запрошена строка №{start_row} и далее.")

    data = []
    for row_obj in rows[start_row - 1:]:
        cells_text = [cell.text.strip() for cell in row_obj.cells]
        data.append(cells_text)
    return data

def flatten_table(list_of_rows, discipline_name=None):
    """
    Превращает список строк (каждая – список ячеек) в плоскую структуру.
    - Строки, где непустые ячейки заканчиваются словом 'семестр' => строка-семестр
    - Строки, где 4-я ячейка 'Тема...' => строка-тема
    - Иначе – обычные данные
    """
    current_semester = None
    current_topic = None
    flattened_data = []

    for row in list_of_rows:
        non_empty = [cell for cell in row if cell]
        # строка-семестр?
        if non_empty:
            unique_vals = set(non_empty)
            if len(unique_vals) == 1 and list(unique_vals)[0].lower().endswith('семестр'):
                current_semester = list(unique_vals)[0]
                continue

        # строка-тема?
        if len(row) >= 4:
            cell_3 = row[3].strip()
            other_cells_empty = all(not row[i].strip() for i in range(len(row)) if i != 3)
            if cell_3.lower().startswith("тема") and other_cells_empty:
                current_topic = cell_3
                continue

        data_dict = {
            'semester':   current_semester,
            'topic':      current_topic,
            'discipline': discipline_name,
        }
        for i, val in enumerate(row, start=1):
            data_dict[f'col{i}'] = val
        flattened_data.append(data_dict)

    return flattened_data

# =============================================================================
# ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ ДЛЯ ОБРАБОТКИ КОЛОНОК
# =============================================================================

def parse_lesson_number(text):
    """ Ищет '№ X/Y', возвращает Y или '' """
    txt = text.replace('\n', ' ')
    pattern = r'№\s*\d+/(\d+)'
    match = re.search(pattern, txt)
    if match:
        return match.group(1).strip()
    return ""

def remove_lesson_number_pattern(text):
    """ Удаляет '№ X/Y' из строки. """
    txt = text.replace('\n', ' ')
    pattern = r'№\s*\S+'
    return re.sub(pattern, '', txt).strip()

def split_first_line(text):
    """
    Разбивает text на (первая_строка, остальное).
    Если нет переносов, остаток = "".
    """
    lines = text.split('\n')
    first_line = lines[0].strip()
    remainder = '\n'.join(line.strip() for line in lines[1:]).strip()
    return first_line, remainder

def remove_any_numbering(text: str) -> str:
    """
    Удаляем ведущие '1.', '1)', '-', '•' и т.п. из каждой строки, без новой нумерации.
    """
    lines = text.split('\n')
    cleaned_lines = []
    for line in lines:
        line = line.strip()
        if not line:
            continue
        # Примеры: "1.", "2)", "• ", "- "
        line = re.sub(r'^\d+[\.)]\s*', '', line)
        line = re.sub(r'^[\-\•]\s*', '', line)
        cleaned_lines.append(line)
    return "\n".join(cleaned_lines)

def expand_number_ranges(cell_value: str) -> str:
    """
    Преобразует '1-4, 5-8' -> '1,2,3,4,5,6,7,8'.
    Если формат не число/диапазон, возвращаем исходное значение.
    """
    if not cell_value:
        return ""

    chunks = cell_value.split(',')
    all_numbers = []

    for chunk in chunks:
        chunk = chunk.strip()
        match = re.match(r'^(\d+)\s*-\s*(\d+)$', chunk)
        if match:
            start = int(match.group(1))
            end   = int(match.group(2))
            if start <= end:
                all_numbers.extend(range(start, end + 1))
            else:
                # если start > end, меням местами
                all_numbers.extend(range(end, start + 1))
        else:
            # просто число?
            if chunk.isdigit():
                all_numbers.append(int(chunk))
            else:
                # неизвестный формат -> возвращаем исходное
                return cell_value

    return ",".join(str(n) for n in all_numbers)

def pick_lines_from_text(source_text: str, lines_str: str) -> str:
    """
    Берет многострочный source_text, разбивает его по \n,
    затем по lines_str (например "2,4") берет 2-ю,4-ю строку (1-based).
    Склеиваем их через \n и возвращаем. Если индекс не в диапазоне, пропускаем.
    """
    if not lines_str.strip():
        return ""

    source_lines = source_text.split('\n')
    indexes_str = lines_str.split(',')
    picked = []

    for idx_str in indexes_str:
        idx_str = idx_str.strip()
        if idx_str.isdigit():
            idx = int(idx_str)
            if 1 <= idx <= len(source_lines):
                picked.append(source_lines[idx-1].strip())

    return "\n".join(picked)

# =============================================================================
# ГЛАВНАЯ ФУНКЦИЯ ПАРСИНГА
# =============================================================================

def parse_docx_to_xlsx(docx_path, xlsx_path):
    """
    1) Открыть DOCX
    2) Извлечь дисциплину
    3) Собрать литературу, мат. обеспечение, знать/уметь/владеть
    4) Считать таблицу => DataFrame
    5) Обработать колонки (семестр, тема, тип занятия, время и т.п.)
    6) col5, col6 => разворачиваем диапазоны "1-4,5" -> "1,2,3,4,5"
    7) Для col5 берем строки из "Материальное обеспечение"
    8) Для col6 берем строки из "Литература"
    9) Переименовываем col5->"материальное обеспечение на занятие"
       col6->"литература на занятие"
    10) Сохраняем в XLSX
    """
    doc = docx.Document(docx_path)

    # (1) Дисциплина
    discipline_name = get_discipline_name(doc)

    # (2) Собираем тексты
    literature_str = parse_text_between(doc, START_REF, END_REF)
    material_str   = parse_text_between(doc, START_MATERIAL, END_MATERIAL)
    know_str       = parse_text_between(doc, KNOW_START, KNOW_END)
    skill_str      = parse_text_between(doc, SKILL_START, SKILL_END)
    master_str     = parse_text_between(doc, MASTER_START, MASTER_END)

    # (3) Чтение таблицы
    table_rows = read_table_from_docx(doc, TABLE_NUMBER, START_ROW)
    flat_data = flatten_table(table_rows, discipline_name=discipline_name)
    df = pd.DataFrame(flat_data)

    # Удаляем col1,col7
    df.drop(columns=[c for c in EXCLUDED_COLS if c in df.columns],
            inplace=True, errors='ignore')

    # (4) Обработка
    # A) "semester" -> цифра
    if 'semester' in df.columns:
        df['semester'] = df['semester'].str.extract(r'(\d+)').fillna('')

    # B) "topic" -> "Номер темы", "Название темы"
    if 'topic' in df.columns:
        pattern_topic = r'^Тема\s+№?\s*(\d+)\.?[\s]*(.*)$'
        extracted = df['topic'].str.extract(pattern_topic, expand=True).fillna('')
        df['Номер темы'] = extracted[0]
        df['Название темы'] = extracted[1]
        df.drop(columns=['topic'], inplace=True)

    # C) col2 => "Тип занятия", "Номер занятия"
    if 'col2' in df.columns:
        df['Номер занятия'] = df['col2'].apply(parse_lesson_number)
        df['col2'] = df['col2'].apply(remove_lesson_number_pattern)
        df.rename(columns={'col2': 'Тип занятия'}, inplace=True)

    # D) col3 => "Время в минутах" (×45)
    if 'col3' in df.columns:
        df.rename(columns={'col3': 'Время в минутах'}, inplace=True)
        df['Время в минутах'] = pd.to_numeric(df['Время в минутах'], errors='coerce').fillna(0)
        df['Время в минутах'] = (df['Время в минутах'] * 45).astype(int)

    # E) col4 => "Учебные вопросы"; первая строка => "Название занятия"
    if 'col4' in df.columns:
        df.rename(columns={'col4': 'Учебные вопросы'}, inplace=True)
        df['Название занятия'], remainder = zip(*df['Учебные вопросы'].apply(split_first_line))
        df['Учебные вопросы'] = remainder

    # F) Добавляем новые поля (литература, матобесп, знать/уметь/владеть)
    df["Литература"] = literature_str
    df["Материальное обеспечение"] = material_str
    df["Знать"] = know_str
    df["Уметь"] = skill_str
    df["Владеть"] = master_str

    # Удаляем нумерацию из "Учебные вопросы"
    if 'Учебные вопросы' in df.columns:
        df['Учебные вопросы'] = df['Учебные вопросы'].apply(remove_any_numbering)

    # --- col5,col6 => разворачиваем диапазоны, потом берём соответствующие строки ---
    if 'col5' in df.columns:
        df['col5'] = df['col5'].apply(expand_number_ranges)
        df['col5'] = df['col5'].apply(lambda s: pick_lines_from_text(material_str, s))

    if 'col6' in df.columns:
        df['col6'] = df['col6'].apply(expand_number_ranges)
        df['col6'] = df['col6'].apply(lambda s: pick_lines_from_text(literature_str, s))

    # Теперь переименовываем col5->"материальное обеспечение на занятие"
    #              и col6->"литература на занятие"
    # (RENAME_MAP уже содержит эти ключи)
    df.rename(columns=RENAME_MAP, inplace=True)

    # Итоговый порядок
    all_cols = list(df.columns)
    final_order = [c for c in DESIRED_ORDER if c in all_cols]
    remaining = [c for c in all_cols if c not in final_order]
    final_order += remaining
    df = df[final_order]

    # (5) Сохраняем
    df.to_excel(xlsx_path, index=False, engine='openpyxl')
    print("Парсинг завершен. Результат сохранен в:", xlsx_path)
    print(df.head(15).to_string(index=False))
    
    return df

# =============================================================================
# ФУНКЦИИ ДЛЯ РАБОТЫ С ШАБЛОНОМ DOCX
# =============================================================================

def replace_placeholders(doc, replacements):
    """
    Заменяет плейсхолдеры в Word документе с учетом особенностей хранения текста в DOCX.
    
    :param doc: объект docx.Document
    :param replacements: словарь {плейсхолдер: значение}
    """
    import tempfile
    import zipfile
    import os
    import re
    
    print("Заменяемые плейсхолдеры:")
    for key, value in replacements.items():
        print(f"  {key} -> {value}")
    
    # Создаем временный файл для документа
    with tempfile.TemporaryDirectory() as tmpdirname:
        # Сохраняем документ во временный файл
        temp_docx = os.path.join(tmpdirname, 'template.docx')
        doc.save(temp_docx)
        
        # Создаем папку для распаковки
        extract_dir = os.path.join(tmpdirname, 'extracted')
        os.makedirs(extract_dir, exist_ok=True)
        
        # Распаковываем docx (это zip-архив)
        with zipfile.ZipFile(temp_docx, 'r') as zip_ref:
            zip_ref.extractall(extract_dir)
        
        # Ищем все XML файлы с содержимым
        xml_path = os.path.join(extract_dir, 'word')
        xml_files = []
        
        # Добавляем основные XML файлы, которые могут содержать текст
        document_xml = os.path.join(xml_path, 'document.xml')
        if os.path.exists(document_xml):
            xml_files.append(document_xml)
            
        # Ищем header*.xml и footer*.xml
        for filename in os.listdir(xml_path):
            if filename.startswith('header') or filename.startswith('footer'):
                if filename.endswith('.xml'):
                    xml_files.append(os.path.join(xml_path, filename))
        
        # Заменяем плейсхолдеры в каждом XML файле
        for xml_file in xml_files:
            try:
                with open(xml_file, 'r', encoding='utf-8') as f:
                    content = f.read()
                
                print(f"\nОбрабатываю файл: {os.path.basename(xml_file)}")
                
                # Подготовим словарь для замены без $ в ключах
                clean_replacements = {}
                for key, value in replacements.items():
                    # Удаляем $ из ключа если он есть
                    clean_key = key
                    if clean_key.startswith('$'):
                        clean_key = clean_key[1:]
                    clean_replacements[clean_key] = value
                
                # Заменяем все плейсхолдеры (с $ и без)
                for key, value in clean_replacements.items():
                    # Пробуем разные варианты плейсхолдеров
                    dollar_key = '$' + key  # С долларом
                    
                    # Замена всех вариантов плейсхолдеров
                    if dollar_key in content:
                        content = content.replace(dollar_key, str(value))
                        print(f"  Заменено: {dollar_key} -> {value}")
                    
                    if key in content:
                        content = content.replace(key, str(value))
                        print(f"  Заменено: {key} -> {value}")
                
                # Сохраняем изменения
                with open(xml_file, 'w', encoding='utf-8') as f:
                    f.write(content)
                    
            except Exception as e:
                print(f"  Ошибка при обработке файла {os.path.basename(xml_file)}: {str(e)}")
                continue
        
        # Создаем обновленный docx файл
        new_docx = os.path.join(tmpdirname, 'result.docx')
        with zipfile.ZipFile(new_docx, 'w') as zipf:
            for root, dirs, files in os.walk(extract_dir):
                for file in files:
                    file_path = os.path.join(root, file)
                    arcname = os.path.relpath(file_path, extract_dir)
                    zipf.write(file_path, arcname)
        
        # Загружаем новый документ
        new_doc = Document(new_docx)
        return new_doc

def generate_lesson_docx(template_path, output_path, lesson_data, form_data):
    """
    Генерирует DOCX файл для занятия, заполняя шаблон данными.
    
    :param template_path: путь к шаблону DOCX
    :param output_path: путь для сохранения результата
    :param lesson_data: DataFrame Series с данными занятия
    :param form_data: словарь с данными из формы
    :return: успешно ли создан документ
    """
    try:
        print(f"\nГенерация документа: {output_path}")
        print(f"Используемый шаблон: {template_path}")
        
        doc = Document(template_path)
        
        # Собираем все замены из формы и данных занятия в один словарь
        replacements = {}
        
        # Сначала добавляем данные из формы (с $)
        for key, value in form_data.items():
            replacements[f"${key}"] = value
        
        # Затем добавляем данные из занятия (lesson_data)
        replacements["$ДИСЦИПЛИНА"] = lesson_data.get("Дисциплина", "")
        replacements["$ВИДЗАНЯТИЯ"] = lesson_data.get("Тип занятия", "")
        replacements["$ТЕМАЗАНЯТИЯ"] = f"№ {lesson_data.get('Номер темы', '')}/{lesson_data.get('Номер занятия', '')} {lesson_data.get('Название темы', '')}"
        
        # Обрабатываем знать/уметь/владеть в зависимости от типа занятия
        lesson_type = lesson_data.get("Тип занятия", "").lower()
        know_text = lesson_data.get("Знать", "")
        skill_text = lesson_data.get("Уметь", "")
        master_text = lesson_data.get("Владеть", "")
        
        # Устанавливаем поля знать/уметь/владеть согласно типу занятия
        if "групповое" in lesson_type:
            replacements["$ЗНАТЬ"] = know_text
            replacements["$УМЕТЬ"] = skill_text
            replacements["$ВЛАДЕТЬ"] = ""
        elif "лекция" in lesson_type:
            replacements["$ЗНАТЬ"] = know_text
            replacements["$УМЕТЬ"] = ""
            replacements["$ВЛАДЕТЬ"] = ""
        elif "практическое" in lesson_type:
            replacements["$ЗНАТЬ"] = ""
            replacements["$УМЕТЬ"] = skill_text
            replacements["$ВЛАДЕТЬ"] = master_text
        elif "семинар" in lesson_type:
            replacements["$ЗНАТЬ"] = know_text
            replacements["$УМЕТЬ"] = ""
            replacements["$ВЛАДЕТЬ"] = ""
        else:
            # По умолчанию
            replacements["$ЗНАТЬ"] = know_text
            replacements["$УМЕТЬ"] = skill_text
            replacements["$ВЛАДЕТЬ"] = master_text
        
        # Форматируем учебные вопросы
        questions = lesson_data.get("Учебные вопросы", "")
        if questions:
            # Форматируем как список с номерами
            lines = questions.strip().split('\n')
            formatted_questions = f"> {lesson_data.get('Название занятия', '')}\n>\n"
            for i, q in enumerate(lines, 1):
                if q.strip():
                    formatted_questions += f"> {i}. {q.strip()}\n"
        else:
            formatted_questions = f"> {lesson_data.get('Название занятия', '')}"
        
        replacements["$УЧЕБНЫЕВОПРОСЫ"] = formatted_questions
        
        # Устанавливаем время, литературу и технические средства
        replacements["$ВРЕМЯ"] = str(lesson_data.get("Время в минутах", ""))
        
        # Используем либо литературу для конкретного занятия, либо общую
        lit_text = lesson_data.get("Литература на занятие", "")
        if not lit_text:
            lit_text = lesson_data.get("Литература", "")
        replacements["$ЛИТЕРАТУРА"] = lit_text
        
        # Технические средства
        replacements["$ТЕХСРЕДСТВА"] = "1. Компьютер\n2. Проектор\n3. Презентация по теме"
        
        print("Сформированы замены для плейсхолдеров:")
        for key, value in replacements.items():
            val_preview = str(value)[:50] + "..." if len(str(value)) > 50 else value
            print(f"  {key} -> {val_preview}")
        
        # Заменяем плейсхолдеры и сохраняем
        new_doc = replace_placeholders(doc, replacements)
        new_doc.save(output_path)
        print(f"Документ успешно сохранен: {output_path}")
        return True
    except Exception as e:
        print(f"Ошибка при создании документа: {str(e)}")
        return False

# Функция сохранения всех занятий как DOCX (для GUI)
def save_all_lessons(parsed_df, template_file, output_dir, form_data):
    """
    Сохраняет все занятия из DataFrame как DOCX файлы
    
    :param parsed_df: DataFrame с данными занятий
    :param template_file: путь к шаблону DOCX
    :param output_dir: директория для сохранения результатов
    :param form_data: словарь с данными из формы
    :return: tuple(количество успешно созданных файлов, общее количество)
    """
    if parsed_df is None or parsed_df.empty:
        return 0, 0
    
    # Счетчики
    total = len(parsed_df)
    success = 0
    
    # Для каждой строки DataFrame создаем файл
    for idx, row in parsed_df.iterrows():
        topic_num = row.get('Номер темы', '')
        lesson_num = row.get('Номер занятия', '')
        
        if not topic_num or not lesson_num:
            continue
            
        # Формируем имя файла
        filename = f"Тема_{topic_num}_Занятие_{lesson_num}.docx"
        output_path = os.path.join(output_dir, filename)
        
        # Генерируем документ
        if generate_lesson_docx(template_file, output_path, row, form_data):
            success += 1
    
    return success, total

# =============================================================================
# GUI: CUSTOMTKINTER
# =============================================================================

def run_gui():
    ctk.set_appearance_mode("System")
    ctk.set_default_color_theme("blue")

    app = ctk.CTk()
    app.title("Парсер DOCX -> XLSX и генератор планов занятий")
    app.geometry("1100x600")

    # Переменные для хранения значений
    selected_docx = tk.StringVar(value="Файл не выбран")
    parsed_data = {"df": None}  # Хранение DataFrame
    
    # Переменные для полей формы
    selected_lesson = tk.StringVar()
    day = tk.StringVar(value="04")
    month = tk.StringVar(value="февраля")
    year = tk.StringVar(value="2025")
    chief = tk.StringVar(value="В. Пупкин")
    group_number = tk.StringVar(value="ИКТВ-11")
    lesson_date = tk.StringVar(value="16.05.2025")
    classroom = tk.StringVar(value="324")
    instructor = tk.StringVar(value="Д. Иванов")
    
    # Создаем основной фрейм
    main_frame = ctk.CTkFrame(app)
    main_frame.pack(fill=tk.BOTH, expand=True)
    
    # Левый фрейм (выбор файла)
    left_frame = ctk.CTkFrame(main_frame)
    left_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=False, padx=10, pady=10)
    
    # Правый фрейм (форма для занятий) - изначально скрыт
    right_frame = ctk.CTkFrame(main_frame)
    
    # Функция выбора DOCX файла
    def choose_docx_file():
        file_path = filedialog.askopenfilename(
            title="Выберите DOCX файл",
            filetypes=[("Word Documents", "*.docx"), ("Все файлы", "*.*")]
        )
        if file_path:
            selected_docx.set(file_path)
            label_docx.configure(text=os.path.basename(file_path))
            
            # Парсим файл
            try:
                # Создаем временный XLSX файл
                temp_dir = os.path.dirname(file_path)
                temp_xlsx = os.path.join(temp_dir, f"temp_{os.path.basename(file_path)}.xlsx")
                
                # Парсим DOCX в XLSX
                parsed_data["df"] = parse_docx_to_xlsx(file_path, temp_xlsx)
                
                # Наполняем выпадающий список занятий
                populate_lesson_dropdown(parsed_data["df"])
                
                # Показываем правый фрейм
                right_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=10, pady=10)
                
                # Удаляем временный файл
                if os.path.exists(temp_xlsx):
                    os.remove(temp_xlsx)
                    
            except Exception as e:
                messagebox.showerror("Ошибка", f"Ошибка при парсинге файла:\n{str(e)}")
    
    # Функция наполнения выпадающего списка занятий
    def populate_lesson_dropdown(df):
        if df is not None and not df.empty:
            # Формируем список занятий в формате "№ topic_num/lesson_num lesson_title"
            lessons = []
            for idx, row in df.iterrows():
                topic_num = row.get('Номер темы', '')
                lesson_num = row.get('Номер занятия', '')
                lesson_title = row.get('Название занятия', '')
                
                if topic_num and lesson_num:
                    formatted = f"№ {topic_num}/{lesson_num} {lesson_title}"
                    lessons.append(formatted)
            
            # Обновляем список в выпадающем меню
            lesson_menu.configure(values=lessons)
            if lessons:
                lesson_menu.set("Выберите занятие")
    
    # Функция обработки файла (оригинальная функциональность)
    def process_file():
        docx_file = selected_docx.get().strip()
        if docx_file == "Файл не выбран" or not os.path.isfile(docx_file):
            messagebox.showwarning("Внимание", "Сначала выберите корректный DOCX файл!")
            return

        xlsx_file = filedialog.asksaveasfilename(
            title="Сохранить XLSX как",
            defaultextension=".xlsx",
            filetypes=[("Excel файлы", "*.xlsx"), ("Все файлы", "*.*")]
        )
        if not xlsx_file:
            return

        try:
            parsed_data["df"] = parse_docx_to_xlsx(docx_file, xlsx_file)
            messagebox.showinfo("Готово", f"Результат сохранён:\n{xlsx_file}")
            
            # Наполняем выпадающий список занятий
            populate_lesson_dropdown(parsed_data["df"])
            
            # Показываем правый фрейм
            right_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=10, pady=10)
            
        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка при парсинге:\n{str(e)}")
    
    # Функция сохранения всех занятий как DOCX
    def save_all_lessons():
        if parsed_data["df"] is None or parsed_data["df"].empty:
            messagebox.showwarning("Внимание", "Сначала загрузите и обработайте DOCX файл!")
            return
        
        output_dir = filedialog.askdirectory(title="Выберите папку для сохранения")
        if not output_dir:
            return
            
        # Получаем шаблон
        template_file = "Template.docx"
        # Проверяем, есть ли файл шаблона в текущей директории
        if not os.path.exists(template_file):
            template_file = filedialog.askopenfilename(
                title="Выберите файл шаблона",
                filetypes=[("Word Documents", "*.docx")]
            )
            if not template_file:
                return
        
        # Форма данных для шаблона
        form_data = {
            "НАЧАЛЬНИК": chief.get(),
            "ЧИСЛА": day.get(),
            "МЕСЯЦА": month.get(),
            "ГОДА": year.get(),
            "ГРУППАНОМЕР": group_number.get(),
            "ДАТАПРОВЕДЕНИЯ": lesson_date.get(),
            "АУДИТОРИЯ": classroom.get(),
            "РУКОВОДИТЕЛЬ": instructor.get()
        }
        
        # Счетчики
        total = len(parsed_data["df"])
        success = 0
        
        # Для каждой строки DataFrame создаем файл
        for idx, row in parsed_data["df"].iterrows():
            topic_num = row.get('Номер темы', '')
            lesson_num = row.get('Номер занятия', '')
            
            if not topic_num or not lesson_num:
                continue
                
            # Формируем имя файла
            filename = f"Тема_{topic_num}_Занятие_{lesson_num}.docx"
            output_path = os.path.join(output_dir, filename)
            
            # Генерируем документ
            if generate_lesson_docx(template_file, output_path, row, form_data):
                success += 1
        
        messagebox.showinfo("Операция завершена", 
                          f"Успешно создано {success} из {total} документов\n"
                          f"Результаты сохранены в:\n{output_dir}")
    
    # Функция сохранения одного выбранного занятия
    def save_single_lesson():
        if parsed_data["df"] is None or parsed_data["df"].empty:
            messagebox.showwarning("Внимание", "Сначала загрузите и обработайте DOCX файл!")
            return
            
        lesson_text = selected_lesson.get()
        if not lesson_text or lesson_text == "Выберите занятие":
            messagebox.showwarning("Внимание", "Выберите занятие из списка!")
            return
            
        # Ищем выбранное занятие в DataFrame
        selected_row = None
        for idx, row in parsed_data["df"].iterrows():
            topic_num = row.get('Номер темы', '')
            lesson_num = row.get('Номер занятия', '')
            lesson_title = row.get('Название занятия', '')
            
            formatted = f"№ {topic_num}/{lesson_num} {lesson_title}"
            if formatted == lesson_text:
                selected_row = row
                break
                
        if selected_row is None:
            messagebox.showerror("Ошибка", "Выбранное занятие не найдено в данных!")
            return
            
        # Выбираем путь сохранения
        output_file = filedialog.asksaveasfilename(
            title="Сохранить как",
            defaultextension=".docx",
            filetypes=[("Word Documents", "*.docx"), ("Все файлы", "*.*")],
            initialfile=f"Тема_{selected_row.get('Номер темы', '')}_Занятие_{selected_row.get('Номер занятия', '')}.docx"
        )
        if not output_file:
            return
            
        # Получаем шаблон
        template_file = "Template.docx"
        # Проверяем, есть ли файл шаблона в текущей директории
        if not os.path.exists(template_file):
            template_file = filedialog.askopenfilename(
                title="Выберите файл шаблона",
                filetypes=[("Word Documents", "*.docx")]
            )
            if not template_file:
                return
                
        # Форма данных для шаблона
        form_data = {
            "НАЧАЛЬНИК": chief.get(),
            "ЧИСЛА": day.get(),
            "МЕСЯЦА": month.get(),
            "ГОДА": year.get(),
            "ГРУППАНОМЕР": group_number.get(),
            "ДАТАПРОВЕДЕНИЯ": lesson_date.get(),
            "АУДИТОРИЯ": classroom.get(),
            "РУКОВОДИТЕЛЬ": instructor.get()
        }
        
        # Генерируем документ
        if generate_lesson_docx(template_file, output_file, selected_row, form_data):
            messagebox.showinfo("Готово", f"Файл успешно создан:\n{output_file}")
        else:
            messagebox.showerror("Ошибка", "Не удалось создать файл!")
            
    # Настройка левого фрейма
    left_title = ctk.CTkLabel(left_frame, text="Выберите DOCX файл:", font=("Arial", 14, "bold"))
    left_title.pack(pady=(15, 5))

    btn_browse = ctk.CTkButton(left_frame, text="Обзор", command=choose_docx_file)
    btn_browse.pack(pady=5)

    label_docx = ctk.CTkLabel(left_frame, text=selected_docx.get(), wraplength=250)
    label_docx.pack(pady=5)

    btn_process = ctk.CTkButton(left_frame, text="Сохранить результат (XLSX)", command=process_file)
    btn_process.pack(pady=(15, 20))
    
    # Настройка правого фрейма
    # Заголовок
    right_title = ctk.CTkLabel(right_frame, text="Генерация плана занятия", font=("Arial", 16, "bold"))
    right_title.pack(pady=(15, 20))
    
    # Создаем фрейм для выбора занятия
    lesson_frame = ctk.CTkFrame(right_frame)
    lesson_frame.pack(fill=tk.X, padx=10, pady=5)
    
    lesson_label = ctk.CTkLabel(lesson_frame, text="Выберите занятие", width=150, anchor="w")
    lesson_label.pack(side=tk.LEFT, padx=5)
    
    lesson_menu = ctk.CTkOptionMenu(lesson_frame, variable=selected_lesson, values=["Выберите занятие"], width=250)
    lesson_menu.pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)
    
    # Дата утверждения
    date_frame = ctk.CTkFrame(right_frame)
    date_frame.pack(fill=tk.X, padx=10, pady=5)
    
    date_label = ctk.CTkLabel(date_frame, text="СОЗДАН (ЧИСЛА), (МЕСЯЦА), (ГОД)", width=150, anchor="w")
    date_label.pack(side=tk.LEFT, padx=5)
    
    day_entry = ctk.CTkEntry(date_frame, textvariable=day, width=60)
    day_entry.pack(side=tk.LEFT, padx=2)
    
    month_entry = ctk.CTkEntry(date_frame, textvariable=month, width=100)
    month_entry.pack(side=tk.LEFT, padx=2)
    
    year_entry = ctk.CTkEntry(date_frame, textvariable=year, width=80)
    year_entry.pack(side=tk.LEFT, padx=2)
    
    # Начальник
    chief_frame = ctk.CTkFrame(right_frame)
    chief_frame.pack(fill=tk.X, padx=10, pady=5)
    
    chief_label = ctk.CTkLabel(chief_frame, text="НАЧАЛЬНИК", width=150, anchor="w")
    chief_label.pack(side=tk.LEFT, padx=5)
    
    chief_entry = ctk.CTkEntry(chief_frame, textvariable=chief, width=250)
    chief_entry.pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)
    
    # Номер группы
    group_frame = ctk.CTkFrame(right_frame)
    group_frame.pack(fill=tk.X, padx=10, pady=5)
    
    group_label = ctk.CTkLabel(group_frame, text="ГРУППАНОМЕР", width=150, anchor="w")
    group_label.pack(side=tk.LEFT, padx=5)
    
    group_entry = ctk.CTkEntry(group_frame, textvariable=group_number, width=150)
    group_entry.pack(side=tk.LEFT, padx=5)
    
    date_lesson_label = ctk.CTkLabel(group_frame, text="ДАТАПРОВЕДЕНИЯ", width=150, anchor="w")
    date_lesson_label.pack(side=tk.LEFT, padx=5)
    
    date_lesson_entry = ctk.CTkEntry(group_frame, textvariable=lesson_date, width=100)
    date_lesson_entry.pack(side=tk.LEFT, padx=5)
    
    # Аудитория и руководитель
    classroom_frame = ctk.CTkFrame(right_frame)
    classroom_frame.pack(fill=tk.X, padx=10, pady=5)
    
    classroom_label = ctk.CTkLabel(classroom_frame, text="АУДИТОРИЯ", width=150, anchor="w")
    classroom_label.pack(side=tk.LEFT, padx=5)
    
    classroom_entry = ctk.CTkEntry(classroom_frame, textvariable=classroom, width=100)
    classroom_entry.pack(side=tk.LEFT, padx=5)
    
    instructor_label = ctk.CTkLabel(classroom_frame, text="РУКОВОДИТЕЛЬ", width=150, anchor="w")
    instructor_label.pack(side=tk.LEFT, padx=5)
    
    instructor_entry = ctk.CTkEntry(classroom_frame, textvariable=instructor, width=150)
    instructor_entry.pack(side=tk.LEFT, padx=5)
    
    # Кнопки генерации файлов
    buttons_frame = ctk.CTkFrame(right_frame)
    buttons_frame.pack(fill=tk.X, padx=10, pady=(20, 5))
    
    save_all_btn = ctk.CTkButton(buttons_frame, text="Сохранить все занятия в DOCX", command=save_all_lessons)
    save_all_btn.pack(side=tk.LEFT, padx=5, expand=True, fill=tk.X)
    
    save_one_btn = ctk.CTkButton(buttons_frame, text="Сохранить одно занятие в DOCX", command=save_single_lesson)
    save_one_btn.pack(side=tk.LEFT, padx=5, expand=True, fill=tk.X)
    
    app.mainloop()


if __name__ == "__main__":
    run_gui()